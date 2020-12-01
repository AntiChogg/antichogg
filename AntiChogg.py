from sys import exit , argv
from PySide2.QtWidgets import QApplication, QMainWindow,  QApplication, QFileDialog, QMessageBox, QDialog
from PySide2.QtCore import QFile
from PySide2.QtUiTools import loadUiType
from docx import Document
from os import path
import pickle
from csv import writer, reader
from bs4 import BeautifulSoup
from googleapiclient.discovery import build


CONFIG_FILE = 'config.pkl'
class MainWindow(QMainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        self.initUi()

    def initUi(self):
        Ui_MainWindow, QtBaseClass = loadUiType('mainWindow.ui')
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        hyperlink_format = '<a href="{link}">{text}</a>'
        label=self.ui.label_3
        urlLink=hyperlink_format.format(link='https://developers.google.com/webmaster-tools/search-console-api/v1/configure', text='Get New')
        label.setOpenExternalLinks(True)
        label.setText(urlLink)
        label=self.ui.label_4
        urlLink=hyperlink_format.format(link='https://www.google.com/cse/create/new', text='Get New')
        label.setOpenExternalLinks(True)
        label.setText(urlLink)
        self.setWindowTitle("Anti Chogg")

        if path.exists(CONFIG_FILE):
            with open(CONFIG_FILE,'rb') as f:
                self.config=pickle.load(f)
                my_api_key=self.config['api']
                my_cse_id=self.config['cse']
        else:
            with open(CONFIG_FILE, 'wb') as f:
                self.config = {'api': my_api_key,'cse': my_cse_id}
                pickle.dump(self.config,f)

        self.ui.apiKey.setText(my_api_key)
        self.ui.cseID.setText(my_cse_id)

        self.ui.findChoggs.clicked.connect(self.findThemChoggers)
        self.ui.loadWord.clicked.connect(self.loadWordFile)
        self.ui.loadChogg.clicked.connect(self.loadChoggFile)
        self.ui.exportHyperlinks.clicked.connect(self.exportList)
        self.ui.exportCSV.clicked.connect(self.exportListCSV)
        self.ui.save.clicked.connect(self.saveQuestions)

    def closeEvent(self,event):
        my_api_key = self.ui.apiKey.text()
        my_cse_id = self.ui.cseID.text()

        with open(CONFIG_FILE, 'wb') as f:
            self.config = {'api': my_api_key,'cse': my_cse_id}
            pickle.dump(self.config,f)
        exit()
        event.accept()

    def google_search(self,search_term, api_key, cse_id, **kwargs):
        service = build("customsearch", "v1", developerKey=api_key)
        res = service.cse().list(q=search_term, cx=cse_id, **kwargs).execute()
        return res['items']

    def findThemChoggers(self):
        my_api_key = self.ui.apiKey.text()
        my_cse_id = self.ui.cseID.text()
        hyperlink_format = '<a href="{link}">{text}</a>'
        for i in range(1,19):
            self.ui.progressBar.setValue(i)
            hyper=1
            question=getattr(self.ui, 'question_'+str(i)).toPlainText()
            print(question)
            question=question.replace('\n', '')
            if  question != '':
                print ('searching')
                try:
                    results=self.google_search(question, my_api_key, my_cse_id, num=2)
                    for j in results:
                        urlLink=hyperlink_format.format(link=j['link'], text=j['link'])
                        label=getattr(self.ui, 'hyperlink'+str(hyper)+'_'+str(i))
                        label.setOpenExternalLinks(True)
                        label.setText(urlLink)
                        hyper+=1
                        print ('result',j['link'])
                except:
                    pass

    def loadWordFile(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getOpenFileName(self,"Load Word File","","Word Files(*.docx )", options=options)
        if fileName !='':
            document = Document(fileName)
            docText = '\n\n'.join(paragraph.text for paragraph in document.paragraphs)
            paragraphs=[]
            for paragraph in document.paragraphs:
                paragraphs.append(paragraph.text)
            paragraphs=[x for x in paragraphs if x !='']
            paragraphs.sort(key=len)
            paragraphs=paragraphs[-18:]
            for i,paragraph in enumerate(paragraphs):
                idx=i+1
                question=getattr(self.ui, 'question_'+str(idx)).setText(paragraph)

    def saveQuestions(self):
        fileName, _ = QFileDialog.getSaveFileName(self,"Save Chogg File")
        if fileName !='':
            saveFileName=fileName+'.chogg'
            with open(saveFileName, 'w', newline='',encoding="utf-8") as file:
                writers = writer(file)
                for i in range(1,19):
                    question=getattr(self.ui, 'question_'+str(i)).toPlainText()
                    writers.writerow([question])

    def loadChoggFile(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getOpenFileName(self,"Load Chogg File","","Chogg Files(*.chogg )", options=options)
        if fileName !='':
            with open(fileName,encoding="utf-8") as f:
                readers = reader(f)
                dataSet = []
                for rowdata in readers:
                    if any(x.strip() for x in rowdata):
                        dataSet.append(rowdata)
                csvdata = list(dataSet)
                print (csvdata)
            for i,paragraph in enumerate(csvdata):
                print(paragraph)
                idx=i+1
                question=getattr(self.ui, 'question_'+str(idx)).setText(paragraph[0])

    def exportList(self):
        fileName, _ = QFileDialog.getSaveFileName(self,"Save Word File")
        myDoc=Document()
        problem=1
        for i in range(1,19):
            if getattr(self.ui, 'hypCheck1_'+str(i)).isChecked():
                hyperlink=getattr(self.ui, 'hyperlink1_'+str(i)).text()
                soup = BeautifulSoup(hyperlink)
                for a in soup.find_all('a', href=True):
                    linked=a['href']
                myDoc.add_paragraph('Problem '+str(problem)+': '+str(linked))
                problem+=1
            if getattr(self.ui, 'hypeCheck2_'+str(i)).isChecked():
                hyperlink=getattr(self.ui, 'hyperlink1_'+str(i)).text()
                soup = BeautifulSoup(hyperlink)
                for a in soup.find_all('a', href=True):
                    linked=a['href']
                myDoc.add_paragraph('Problem '+str(problem)+': '+str(linked))
                problem+=1

        myDoc.save(fileName+'.docx')

    def exportListCSV(self):
        fileName, _ = QFileDialog.getSaveFileName(self,"Save CSV File")
        with open(fileName+'.csv') as csv_file:
            chogg_writer = writer(csv_file, delimiter=',')
            for i in range(1,19):
                if getattr(self.ui, 'hypCheck1_'+str(i)).isChecked():
                    hyperlink=getattr(self.ui, 'hyperlink1_'+str(i)).text()
                    soup = BeautifulSoup(hyperlink)
                    for a in soup.find_all('a', href=True):
                        linked=a['href']
                    chogg_writer([str(linked)])
                if getattr(self.ui, 'hypeCheck2_'+str(i)).isChecked():
                    hyperlink=getattr(self.ui, 'hyperlink1_'+str(i)).text()
                    soup = BeautifulSoup(hyperlink)
                    for a in soup.find_all('a', href=True):
                        linked=a['href']
                    chogg_writer([str(linked)])

if __name__ == "__main__":
    app = QApplication(argv)
    window = MainWindow()
    window.show()
    exit(app.exec_())
